import streamlit as st
import ezdxf
from ezdxf.enums import TextEntityAlignment
import pandas as pd
import plotly.graph_objects as go
import io
import datetime
import math
import random
import os
import zipfile
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from openai import OpenAI

# --- 0. 图表全局防乱码配置 ---
font_path_cloud = "AlibabaPuHuiTi-3-115-Black.ttf"
if os.path.exists(font_path_cloud): prop = fm.FontProperties(fname=font_path_cloud)
else: prop = fm.FontProperties(family='SimHei')
plt.rcParams['font.family'] = prop.get_name()
plt.rcParams['axes.unicode_minus'] = False

# --- 1. 配置与 AI 初始化 ---
try: api_key = st.secrets["DEEPSEEK_API_KEY"]
except: api_key = None
if not api_key:
    st.error("⚠️ 架构师警告：本地未检测到 API Key！\n请配置 `.streamlit/secrets.toml`。")
    st.stop()
client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")

def simulate_macro_economy():
    month = datetime.datetime.now().month
    return int(3800 + math.sin(month * math.pi / 6) * 150 + random.uniform(-30, 30)), int(68000 + random.uniform(-400, 400)), round(95 + month * 0.5 + random.uniform(-1, 1), 1)

GEO_MODELS = {
    "广东 (高地价/高人工)": {"factor": 1.3, "power_price": 0.75, "desc": "需追求极致自动化与高容积率。"},
    "内蒙古 (低电价/严寒)": {"factor": 0.85, "power_price": 0.35, "desc": "地价便宜可铺开建设，需重度保温防冻。"},
    "四川 (丰水期/水电/高湿)": {"factor": 1.0, "power_price": 0.55, "desc": "常年高湿，车间必须增加大功率除湿。"}
}
DEFECT_MODELS = {"饼干/烘焙": "【烤焦边缘】", "乳制品/液态奶": "【包装封口漏气】", "肉制品/切片": "【真空袋漏气】"}

def calculate_physics(cap, cat, province):
    f = {"饼干/烘焙": 1.0, "乳制品/液态奶": 1.5, "肉制品/切片": 1.3}.get(cat, 1.0)
    geo = GEO_MODELS[province]
    return {"elec": round(cap * 15 * f, 1), "water": round(cap * 3.5 * f, 1), "steam": round(cap * 0.8 * f, 1), "cost_factor": geo["factor"], "power_price": geo["power_price"], "geo_desc": geo["desc"]}

# --- 2. 满血全量 BOM (彻底修复：100%找回所有专业、数量与智能运维建议！) ---
@st.cache_data
def generate_detailed_bom(category, capacity, cost_factor, province):
    defect_target = DEFECT_MODELS[category]
    base_bom = [
        {"专业": "给排水", "设备与材料": "RO反渗透与超滤水站", "规格参数": "直饮级", "数量": 1, "预估造价": round(35 * cost_factor, 1), "合规说明": "满足食品接触水标准", "智能运维建议": "每日在线检测水质硬度与余氯"},
        {"专业": "环保", "设备与材料": "工业污水处理站", "规格参数": "A/O生化工艺", "数量": 1, "预估造价": round(45 * cost_factor, 1), "合规说明": "下风向隔离布置", "智能运维建议": "COD与氨氮指标实时联网上传环保局"},
        {"专业": "消防", "设备与材料": "环形消防管网与消防泵房", "规格参数": "≥4米环形通道", "数量": 1, "预估造价": round(25 * cost_factor, 1), "合规说明": "消防车道绝对畅通", "智能运维建议": "消防栓水压传感器实时联动火灾报警主板"},
        {"专业": "智能化", "设备与材料": "AI 视觉次品剔除站", "规格参数": "边缘算力盒+高速相机", "数量": 1, "预估造价": 45.0, "合规说明": f"拦截: {defect_target}", "智能运维建议": "先进行数据增强处理，再定期补充负样本微调模型，防止数据漂移"},
        {"专业": "卫生(SC标准)", "设备与材料": "一更/二更/风淋强制隔离舱", "规格参数": "门禁互锁", "数量": 1, "预估造价": round(20 * cost_factor, 1), "合规说明": "包装区唯一合法入口", "智能运维建议": "风淋室HEPA初中效滤网压差自动报警更换"},
        {"专业": "物流", "设备与材料": "自动码垛机械臂与AGV", "规格参数": "智能仓储调度", "数量": 1, "预估造价": 65.0, "合规说明": "取代人工作业防污染", "智能运维建议": "AGV电量实时监控，设定闲时自动回充策略"}
    ]
    if category == "饼干/烘焙": spec = [{"专业": "工艺", "设备与材料": "燃气隧道炉及成形线", "规格参数": f"{capacity}T", "数量": 1, "预估造价": round(185*cost_factor,1), "合规说明": "排烟罩需伸出设备40cm", "智能运维建议": "动态监测各温区PID炉温曲线"}]
    elif category == "乳制品/液态奶": spec = [{"专业": "工艺", "设备与材料": "UHT杀菌机与无菌灌装", "规格参数": "百级净化", "数量": 1, "预估造价": round(280*cost_factor,1), "合规说明": "蒸汽需配置疏水阀排冷凝水", "智能运维建议": "严密监控137℃高温瞬时保温时间"}]
    else: spec = [{"专业": "工艺", "设备与材料": "深冷速冻库与真空包装机", "规格参数": "-35℃", "数量": 1, "预估造价": round(220*cost_factor,1), "合规说明": "生熟区清洗水池严格分开", "智能运维建议": "根据结霜厚度智能调节化霜周期"}]
    return pd.DataFrame(base_bom + spec)

# --- 3. 终极参数化CAD (新增：标准图框与结构轴网气泡标注) ---
def create_professional_dxf(capacity, category, target_area, doc_type="总图"):
    doc = ezdxf.new('R2010'); doc.styles.new('CHS', dxfattribs={'font': 'simsun.ttc'})
    doc.layers.add("BUILDING", color=7); doc.layers.add("TECH_EQUIP", color=3)
    doc.layers.add("PLUMBING_DRAIN", color=5); doc.layers.add("HVAC_AIR", color=1)
    doc.layers.add("DIMENSION", color=8); doc.layers.add("PROPERTY_LINE", color=1)
    doc.layers.add("FIRE_LANE", color=1, linetype="DASHDOT"); doc.layers.add("GREENERY", color=3)
    doc.layers.add("STRUCT_GRID", color=8, linetype="DASHDOT"); doc.layers.add("STRUCT_COLUMN", color=4)
    doc.layers.add("MEP_MAIN", color=2, lineweight=35); doc.layers.add("HVAC_DUCT", color=6, lineweight=25)
    doc.layers.add("BUILDING_LOG", color=6); doc.layers.add("BUILDING_LIFE", color=7)
    doc.layers.add("TITLE_BLOCK", color=7) # 新增：标准图框专用图层
    msp = doc.modelspace()
    
    sf = math.sqrt(target_area / 8000.0) 
    
    # 智能尺寸标注绘图器
    def draw_building(name, x, y, w, h, layer, color, dims=['B', 'L']):
        rx, ry, rw, rh = x*sf, y*sf, w*sf, h*sf
        msp.add_lwpolyline([(rx, ry), (rx+rw, ry), (rx+rw, ry+rh), (rx, ry+rh), (rx, ry)], dxfattribs={'layer': layer, 'color': color})
        if doc_type in ["总图", "工艺"]: msp.add_text(name, dxfattribs={'height': 150*sf, 'style': 'CHS', 'color': color}).set_placement((rx+rw/2, ry+rh/2), align=TextEntityAlignment.MIDDLE_CENTER)
        
        dim_gap = 600 * sf; text_h = 120 * sf
        real_w = round(rw * 10 / 1000, 1); real_h = round(rh * 10 / 1000, 1)
        if 'B' in dims:
            msp.add_line((rx, ry-dim_gap), (rx+rw, ry-dim_gap), dxfattribs={'layer': 'DIMENSION'}); msp.add_line((rx, ry), (rx, ry-dim_gap-200*sf), dxfattribs={'layer': 'DIMENSION'}); msp.add_line((rx+rw, ry), (rx+rw, ry-dim_gap-200*sf), dxfattribs={'layer': 'DIMENSION'})
            msp.add_text(f"{real_w}m", dxfattribs={'height': text_h, 'style': 'CHS', 'color': 8}).set_placement((rx+rw/2, ry-dim_gap+100*sf), align=TextEntityAlignment.BOTTOM_CENTER)
        if 'T' in dims:
            msp.add_line((rx, ry+rh+dim_gap), (rx+rw, ry+rh+dim_gap), dxfattribs={'layer': 'DIMENSION'}); msp.add_line((rx, ry+rh), (rx, ry+rh+dim_gap+200*sf), dxfattribs={'layer': 'DIMENSION'}); msp.add_line((rx+rw, ry+rh), (rx+rw, ry+rh+dim_gap+200*sf), dxfattribs={'layer': 'DIMENSION'})
            msp.add_text(f"{real_w}m", dxfattribs={'height': text_h, 'style': 'CHS', 'color': 8}).set_placement((rx+rw/2, ry+rh+dim_gap+100*sf), align=TextEntityAlignment.BOTTOM_CENTER)
        if 'L' in dims:
            msp.add_line((rx-dim_gap, ry), (rx-dim_gap, ry+rh), dxfattribs={'layer': 'DIMENSION'}); msp.add_line((rx, ry), (rx-dim_gap-200*sf, ry), dxfattribs={'layer': 'DIMENSION'}); msp.add_line((rx, ry+rh), (rx-dim_gap-200*sf, ry+rh), dxfattribs={'layer': 'DIMENSION'})
            txt = msp.add_text(f"{real_h}m", dxfattribs={'height': text_h, 'style': 'CHS', 'color': 8})
            txt.set_placement((rx-dim_gap-100*sf, ry+rh/2), align=TextEntityAlignment.MIDDLE_CENTER); txt.dxf.rotation = 90
        if 'R' in dims:
            msp.add_line((rx+rw+dim_gap, ry), (rx+rw+dim_gap, ry+rh), dxfattribs={'layer': 'DIMENSION'}); msp.add_line((rx+rw, ry), (rx+rw+dim_gap+200*sf, ry), dxfattribs={'layer': 'DIMENSION'}); msp.add_line((rx+rw, ry+rh), (rx+rw+dim_gap+200*sf, ry+rh), dxfattribs={'layer': 'DIMENSION'})
            txt = msp.add_text(f"{real_h}m", dxfattribs={'height': text_h, 'style': 'CHS', 'color': 8})
            txt.set_placement((rx+rw+dim_gap+100*sf, ry+rh/2), align=TextEntityAlignment.MIDDLE_CENTER); txt.dxf.rotation = 90

    # 1. 全局红线与标准制图标题
    bound_min_x, bound_min_y = -6000 * sf, -7000 * sf
    bound_max_x, bound_max_y = 23000 * sf, 11000 * sf
    msp.add_lwpolyline([(bound_min_x, bound_min_y), (bound_max_x, bound_min_y), (bound_max_x, bound_max_y), (bound_min_x, bound_max_y), (bound_min_x, bound_min_y)], dxfattribs={'layer': 'PROPERTY_LINE', 'lineweight': 50})
    msp.add_text(f"项目规划地界红线 (占地约束: {target_area} ㎡)", dxfattribs={'height': 250*sf, 'style': 'CHS', 'color': 1}).set_placement((0, bound_max_y - 800*sf))
    msp.add_text(f"【{doc_type}】 {category} 工业园专业规划图", dxfattribs={'height': 350*sf, 'style': 'CHS', 'color': 7}).set_placement((8500*sf, bound_max_y + 1000*sf), align=TextEntityAlignment.MIDDLE_CENTER)

    # 【核心新增：工程标准图框/标题栏 (Title Block)】
    tb_w, tb_h = 7000 * sf, 1500 * sf
    tb_x, tb_y = bound_max_x - tb_w, bound_min_y
    msp.add_lwpolyline([(tb_x, tb_y), (tb_x+tb_w, tb_y), (tb_x+tb_w, tb_y+tb_h), (tb_x, tb_y+tb_h), (tb_x, tb_y)], dxfattribs={'layer': 'TITLE_BLOCK', 'color': 7})
    msp.add_line((tb_x, tb_y + tb_h/2), (tb_x+tb_w, tb_y + tb_h/2), dxfattribs={'layer': 'TITLE_BLOCK'})
    msp.add_text(f"工程名称: {category}数字化示范工厂", dxfattribs={'height': 150*sf, 'style': 'CHS', 'color': 7}).set_placement((tb_x + 300*sf, tb_y + tb_h*0.75), align=TextEntityAlignment.MIDDLE_LEFT)
    msp.add_text(f"图纸类别: {doc_type}   |   制图日期: {datetime.datetime.now().strftime('%Y-%m-%d')}", dxfattribs={'height': 150*sf, 'style': 'CHS', 'color': 7}).set_placement((tb_x + 300*sf, tb_y + tb_h*0.25), align=TextEntityAlignment.MIDDLE_LEFT)

    # 2. 基础建筑与全域尺寸标注
    py = 1000; ph = 3500; end_x = 15000; robot_x = end_x + 1500
    draw_building("前处理区", 0, py, 3500, ph, "TECH_EQUIP", 3, ['B', 'L'])
    draw_building("核心热加工区", 4000, py, 7000, ph, "TECH_EQUIP", 3, ['B'])
    draw_building("包装正压区", 11500, py, 3500, ph, "TECH_EQUIP", 3, ['B', 'R'])
    draw_building("原辅料立体库", -4000, 6500, 6000, 3000, "BUILDING_LOG", 6, ['T', 'L'])
    draw_building("成品发货月台", 16500, 6500, 5000, 3000, "BUILDING_LOG", 6, ['T', 'R'])

    # ---------------- 结构图专属 (带轴号气泡) ----------------
    if doc_type == "结构":
        grid_x_idx = 1
        for x in range(0, int(end_x*sf)+1, int(8000*sf)):
            msp.add_line((x, (py-500)*sf), (x, (py+ph+500)*sf), dxfattribs={'layer': 'STRUCT_GRID'})
            # 【核心新增：底部 X 轴编号气泡】
            msp.add_circle((x, (py-800)*sf), radius=200*sf, dxfattribs={'layer': 'STRUCT_GRID'})
            msp.add_text(str(grid_x_idx), dxfattribs={'height': 150*sf, 'style': 'CHS', 'color': 8}).set_placement((x, (py-800)*sf), align=TextEntityAlignment.MIDDLE_CENTER)
            grid_x_idx += 1
            for y in [py*sf, (py+ph/2)*sf, (py+ph)*sf]:
                cw = 600 * sf
                msp.add_lwpolyline([(x-cw/2, y-cw/2), (x+cw/2, y-cw/2), (x+cw/2, y+cw/2), (x-cw/2, y+cw/2), (x-cw/2, y-cw/2)], dxfattribs={'layer': 'STRUCT_COLUMN'})
        # 【核心新增：左侧 Y 轴字母气泡】
        grid_y_idx = 'A'
        for y in [py*sf, (py+ph/2)*sf, (py+ph)*sf]:
            msp.add_line(((-500)*sf, y), ((end_x+500)*sf, y), dxfattribs={'layer': 'STRUCT_GRID'})
            msp.add_circle(((-800)*sf, y), radius=200*sf, dxfattribs={'layer': 'STRUCT_GRID'})
            msp.add_text(grid_y_idx, dxfattribs={'height': 150*sf, 'style': 'CHS', 'color': 8}).set_placement(((-800)*sf, y), align=TextEntityAlignment.MIDDLE_CENTER)
            grid_y_idx = chr(ord(grid_y_idx) + 1)

    # ---------------- 机电图专属 ----------------
    elif doc_type == "机电":
        draw_building("⚡ 高压配电中心", 4000, py-2000, 3000, 1500, "MEP_MAIN", 2, ['B', 'L'])
        draw_building("💧 纯水制备站", 7500, py-2000, 2000, 1500, "PLUMBING_DRAIN", 5, ['B'])
        msp.add_line((-4000*sf, (py+ph/2)*sf), (21000*sf, (py+ph/2)*sf), dxfattribs={'layer': 'MEP_MAIN'})
        msp.add_text("=== 380V 主强电电缆桥架 ===", dxfattribs={'height': 150*sf, 'style': 'CHS', 'color': 2}).set_placement((6000*sf, (py+ph/2+300)*sf))

    # ---------------- 暖通图专属 ----------------
    elif doc_type == "暖通":
        draw_building("❄️ 组合式空调机组", 11500, py+ph+800, 3500, 1500, "HVAC_DUCT", 6, ['T', 'R'])
        msp.add_line((13250*sf, (py+ph+800)*sf), (13250*sf, py*sf), dxfattribs={'layer': 'HVAC_DUCT'})
        msp.add_text("↑ 百级送风主管", dxfattribs={'height': 120*sf, 'style': 'CHS', 'color': 6}).set_placement((13600*sf, (py+ph/2)*sf), align=TextEntityAlignment.MIDDLE_LEFT)

    # ---------------- 满血版总图专属 (全量生态恢复) ----------------
    else: 
        draw_building("一更->二更->风淋室", 10000, py-1500, 1500, 1500, "BUILDING_LIFE", 7, ['B', 'L'])
        draw_building("污水处理站(下风向)", 14000, -5000, 4000, 3000, "BUILDING", 4, ['B', 'R'])
        draw_building("动力中心", 4000, py-2000, 3000, 1500, "BUILDING_LIFE", 7, ['B'])
        draw_building("🌿 厂区景观绿化", -4000, 2500, 3000, 3000, "GREENERY", 3, ['L'])
        draw_building("🌿 防尘绿化带", -4000, -2500, 3000, 2500, "GREENERY", 3, ['L'])
        draw_building("综合办公楼", -4000, -5000, 3000, 1500, "BUILDING_LIFE", 7, ['B', 'L'])
        draw_building("员工食堂", -500, -5000, 2500, 1500, "BUILDING_LIFE", 7, ['B'])
        draw_building("倒班宿舍楼", 2500, -5000, 2000, 1500, "BUILDING_LIFE", 7, ['B'])
        draw_building("独立更衣/淋浴室", 5000, -5000, 3000, 1500, "BUILDING_LIFE", 7, ['B'])
        
        msp.add_line((0, (py-800)*sf), (16000*sf, (py-800)*sf), dxfattribs={'layer': 'PLUMBING_DRAIN', 'lineweight': 35})
        msp.add_text("💧 主排污地沟(1.5%坡度) -> 流向污水站", dxfattribs={'height': 120*sf, 'style': 'CHS', 'color': 5}).set_placement((2000*sf, (py-500)*sf))
        msp.add_line((-4000*sf, (py-2500)*sf), (16000*sf, (py-2500)*sf), dxfattribs={'layer': 'ELEC_POWER', 'lineweight': 25})
        msp.add_text("⚡ 380V 主桥架 / ♨️ 蒸汽管廊", dxfattribs={'height': 120*sf, 'style': 'CHS', 'color': 2}).set_placement((0, (py-2200)*sf))

        msp.add_lwpolyline([(-1000*sf, -1000*sf), (16000*sf, -1000*sf), (16000*sf, 5500*sf), (-1000*sf, 5500*sf), (-1000*sf, -1000*sf)], dxfattribs={'layer': 'FIRE_LANE', 'lineweight': 30})
        msp.add_text("🚒 ≥4米 环形消防通道", dxfattribs={'height': 150*sf, 'style': 'CHS', 'color': 1}).set_placement((7500*sf, -700*sf), align=TextEntityAlignment.MIDDLE_CENTER)
        
        msp.add_line((-4000*sf, 6000*sf), (21000*sf, 6000*sf), dxfattribs={'layer': 'BUILDING_LOG', 'color': 3, 'linetype': 'DASHED'})
        msp.add_text("🤖 AGV 自动物流接驳线", dxfattribs={'height': 150*sf, 'style': 'CHS', 'color': 3}).set_placement((7500*sf, 6200*sf))
        
        msp.add_circle((robot_x*sf, (py + 1700)*sf), radius=1500*sf, dxfattribs={'layer': 'BUILDING_LOG', 'color': 2})
        msp.add_text("🤖码垛机械臂", dxfattribs={'height': 120*sf, 'style': 'CHS', 'color': 2}).set_placement((robot_x*sf, (py+3600)*sf), align=TextEntityAlignment.MIDDLE_CENTER)
        ai_x = 11200; ai_y = py + ph/2
        msp.add_circle((ai_x*sf, ai_y*sf), radius=300*sf, dxfattribs={'layer': 'TECH_EQUIP', 'color': 6})
        msp.add_text(f"📷 视觉拦截: {DEFECT_MODELS[category].split('】')[0].replace('【', '')}", dxfattribs={'height': 120*sf, 'style': 'CHS', 'color': 6}).set_placement((ai_x*sf, (ai_y+500)*sf), align=TextEntityAlignment.MIDDLE_CENTER)

    buf = io.StringIO(); doc.write(buf); return buf.getvalue()

@st.cache_data
def generate_all_dxf_zip(capacity, category, target_area):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("01_建筑总平规划图(图框标准版).dxf", create_professional_dxf(capacity, category, target_area, "总图"))
        zf.writestr("02_土建结构柱网布置图(带轴号).dxf", create_professional_dxf(capacity, category, target_area, "结构"))
        zf.writestr("03_机电(强电弱电给水)综合图.dxf", create_professional_dxf(capacity, category, target_area, "机电"))
        zf.writestr("04_暖通(HVAC)气流压差控制图.dxf", create_professional_dxf(capacity, category, target_area, "暖通"))
    return zip_buffer.getvalue()

# --- 4. Plotly 3D 引擎与 Word 文档 ---
@st.cache_data 
def create_3d_factory(category, target_area):
    fig = go.Figure()
    sf = math.sqrt(target_area / 8000.0)
    def add_3d_building(x, y, z, dx, dy, dz, name, color):
        rx, ry, rz, rdx, rdy, rdz = x*sf, y*sf, z*sf, dx*sf, dy*sf, dz*sf
        xx = [rx, rx+rdx, rx+rdx, rx, rx, rx+rdx, rx+rdx, rx]; yy = [ry, ry, ry+rdy, ry+rdy, ry, ry, ry+rdy, ry+rdy]; zz = [rz, rz, rz, rz, rz+rdz, rz+rdz, rz+rdz, rz+rdz]
        i = [0, 0, 4, 4, 0, 0, 2, 2, 0, 0, 1, 1]; j = [1, 2, 5, 6, 1, 5, 3, 7, 3, 7, 2, 6]; k = [2, 3, 6, 7, 5, 4, 7, 6, 7, 4, 6, 5]
        fig.add_trace(go.Mesh3d(x=xx, y=yy, z=zz, i=i, j=j, k=k, color=color, opacity=0.8, name=name, showlegend=True, hoverinfo="name"))
    
    add_3d_building(0, 10, 0, 35, 35, 8, "前处理区", "#2ca02c"); add_3d_building(40, 10, 0, 70, 35, 8, "热加工区", "#98df8a")
    add_3d_building(115, 10, 0, 35, 35, 8, "包装区", "#2ca02c"); add_3d_building(100, -5, 0, 15, 15, 4, "SC 缓冲舱", "#c7c7c7")
    add_3d_building(40, -5, 0, 30, 15, 6, "动力中心", "#7f7f7f"); add_3d_building(-40, 65, 0, 60, 30, 20, "立体原辅料库", "#1f77b4")
    add_3d_building(165, 65, 0, 50, 30, 12, "发货月台", "#aec7e8"); add_3d_building(140, -50, 0, 40, 30, 5, "污水处理站", "#17becf")
    add_3d_building(-40, -50, 0, 30, 15, 12, "办公区", "#7f7f7f"); add_3d_building(-5, -50, 0, 25, 15, 8, "食堂", "#c7c7c7")

    fig.add_trace(go.Scatter3d(x=[-10*sf, 160*sf, 160*sf, -10*sf, -10*sf], y=[-10*sf, -10*sf, 55*sf, 55*sf, -10*sf], z=[0.5, 0.5, 0.5, 0.5, 0.5], mode='lines', line=dict(color='red', width=4), name="消防通道"))
    fig.add_trace(go.Scatter3d(x=[112*sf], y=[27.5*sf], z=[4*sf], mode='markers+text', marker=dict(size=12, color='red', symbol='cross'), text=["📷 AI视觉"], textposition="top center", name="AI 品控节点", textfont=dict(color='red', size=12)))
    fig.update_layout(scene=dict(xaxis=dict(showgrid=False), yaxis=dict(showgrid=False), zaxis=dict(showgrid=False), aspectmode='data', camera=dict(eye=dict(x=1.2, y=-1.8, z=1.0))), margin=dict(l=0, r=0, b=0, t=30), height=500, legend=dict(x=0, y=1, bgcolor="rgba(255,255,255,0.5)"))
    return fig

def generate_word_manual(project_name, province, category, capacity, ai_long_content, defect_target, area):
    doc = Document(); style = doc.styles['Normal']; style.font.name = '宋体'; style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); style.font.size = Pt(11)
    doc.add_heading(f"{project_name}\n甲级设计院规范书", 0).alignment = 1 
    for p in ai_long_content.split('\n'):
        if p.strip() == "": continue
        if p.startswith('###') or p.startswith('**'): doc.add_heading(p.replace('#', '').replace('*', '').strip(), level=2)
        else: doc.add_paragraph(p.replace('*', '').strip())
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

# --- 5. 界面构建 ---
st.set_page_config(page_title="大一统无损交付基座", layout="wide")
steel_p, copper_p, carbon_p = simulate_macro_economy()

with st.sidebar:
    st.header("🗂️ 总工规划沙盘")
    proj_name = st.text_input("项目名称", "2026 甲级标准工厂")
    prov = st.selectbox("拟建区域", ["广东 (高地价/高人工)", "内蒙古 (低电价/严寒)", "四川 (丰水期/水电/高湿)"])
    cat = st.selectbox("食品门类", ["饼干/烘焙", "乳制品/液态奶", "肉制品/切片"])
    cap = st.slider("日产吨数", 5, 50, 20)
    target_area = st.number_input("实际地块面积(红线约束)", min_value=2000, max_value=20000, value=8000, step=500)
    st.divider()

physics = calculate_physics(cap, cat, prov)
bom_df = generate_detailed_bom(cat, cap, physics['cost_factor'], prov)

st.title("🏭 大一统终极版：BOM 满血回归 × CAD 标题栏标准")
st.plotly_chart(create_3d_factory(cat, target_area), use_container_width=True)

st.divider(); st.subheader("📥 一次性完美交付包 (No Compromise)")
e1, e2 = st.columns(2)
excel_buf = io.BytesIO()
with pd.ExcelWriter(excel_buf, engine='openpyxl') as writer: bom_df.to_excel(writer, index=False)
# 彻底修复：BOM清单一键下载，确保包含“智能运维建议”
e1.download_button("📊 下载满血版 BOM (含全部智能运维与数量)", excel_buf.getvalue(), "Perfect_BOM.xlsx", use_container_width=True)

zip_bytes = generate_all_dxf_zip(cap, cat, target_area)
# 彻底升级：带有图框和轴号的顶级 CAD 矩阵
e2.download_button("🗜️ 下载标准化图纸矩阵 (带图框与轴号)", zip_bytes, "Standardized_Drawings.zip", mime="application/zip", use_container_width=True)

if st.button("🚀 流式生成《总规评审与 AI 模型建议书》", use_container_width=True):
    prompt = f"以30年经验工程总工身份撰写。针对【{prov}】项目，说明在 {target_area}平米限制下容积率达标情况，以及 AI视觉拦截{DEFECT_MODELS[cat]}的部署策略（重点说明针对此类缺陷如何补充负样本）。1500字。"
    st.markdown("### 🤖 总工正在流式打印建议书...")
    report_container = st.empty()
    full_content = ""
    try:
        for chunk in client.chat.completions.create(model="deepseek-chat", messages=[{"role": "user", "content": prompt}], stream=True):
            if chunk.choices[0].delta.content:
                full_content += chunk.choices[0].delta.content
                report_container.markdown(full_content + " ▌")
        report_container.markdown(full_content)
        st.session_state['word_content'] = full_content
        st.success("✅ 生成完毕！")
    except Exception as e: st.error(f"失败: {e}")

if 'word_content' in st.session_state: 
    st.download_button("📄 下载总工审查建议书 (Word)", generate_word_manual(proj_name, prov, cat, cap, st.session_state['word_content'], DEFECT_MODELS[cat], target_area), "Perfect_Proposal.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)